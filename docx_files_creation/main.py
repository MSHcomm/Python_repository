# This Python script to create MSH.docx file. Also, read text form CV21.docx.
# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
import docx

word_document = docx.Document()

word_document.add_heading(" MSH Telecom Engineer \n", 0)

word_document.add_paragraph(' I am locking for join a good reputation company ', 0)
# p = docx.add_paragraph(" I am locking for join a good reputation company \\n", 0)
'''p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True


document.add_paragraph(
    'my current job: Control System Engineer', style='List Bullet'
)
document.add_paragraph(
    'my last job: Embedded Software Engineer', style='List Number'
)

'''
word_document.add_page_break()

word_document.save('MSH.docx')